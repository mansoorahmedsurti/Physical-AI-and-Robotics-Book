import asyncio
import sys
import os
from pathlib import Path
import hashlib
import PyPDF2
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
import numpy as np
import cohere
from qdrant_client import QdrantClient
from qdrant_client.http import models
import asyncpg
from typing import List, Dict, Any
import tempfile


class StandaloneDocumentProcessor:
    def __init__(self, cohere_api_key: str, qdrant_host: str, qdrant_api_key: str):
        self.co = cohere.Client(cohere_api_key)
        self.qdrant_client = QdrantClient(url=qdrant_host, api_key=qdrant_api_key)
        self.collection_name = "document_chunks"

    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks"""
        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # If we're near the end, extend to include the rest
            if end > len(text):
                end = len(text)

            chunk_text = text[start:end]
            chunks.append({
                'content': chunk_text,
                'start_pos': start,
                'end_pos': end
            })

            # Move start position forward, accounting for overlap
            start = end - overlap

            # If the next chunk would be empty, break
            if start >= len(text):
                break

        return chunks

    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from a PDF file"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error extracting text from PDF {pdf_path}: {str(e)}")
            raise
        return text

    def _extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from a DOCX file"""
        try:
            doc = DocxDocument(docx_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)

            # Also extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)

            return '\n'.join(text)
        except Exception as e:
            print(f"Error extracting text from DOCX {docx_path}: {str(e)}")
            raise

    def _extract_text_from_image(self, image_path: str) -> str:
        """Extract text from an image file using OCR"""
        try:
            image = Image.open(image_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            print(f"Error extracting text from image {image_path}: {str(e)}")
            raise

    async def generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text using Cohere"""
        try:
            response = self.co.embed(
                texts=[text],
                model="embed-multilingual-v3.0",  # Using multilingual model for broader language support
                input_type="search_document"  # Required parameter for the v3 model
            )

            # Extract the first (and only) embedding from the response
            return response.embeddings[0]
        except Exception as e:
            print(f"Error generating embedding: {str(e)}")
            raise

    async def process_file(self, file_path: str) -> List[Dict[str, Any]]:
        """Process a file from disk and return chunks with embeddings"""
        try:
            file_extension = Path(file_path).suffix.lower()

            if file_extension == '.pdf':
                content = self._extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                content = self._extract_text_from_docx(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp']:
                content = self._extract_text_from_image(file_path)
            elif file_extension in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv']:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
            else:
                # For other files, try to read as text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()

            # Split content into chunks
            chunks = self._chunk_text(content)

            # Generate embeddings for each chunk
            processed_chunks = []
            for i, chunk in enumerate(chunks):
                embedding = await self.generate_embedding(chunk['content'])

                processed_chunk = {
                    'content': chunk['content'],
                    'embedding': embedding,
                    'page_number': 1  # Simplified for this standalone processor
                }
                processed_chunks.append(processed_chunk)

            return processed_chunks
        except Exception as e:
            print(f"Error processing file {file_path}: {str(e)}")
            raise

    def store_embeddings(self, document_id: str, chunks: List[Dict[str, Any]]) -> List[str]:
        """Store document chunks with their embeddings in Qdrant"""
        points = []
        for idx, chunk in enumerate(chunks):
            point = models.PointStruct(
                id=f"{document_id}_{idx}",
                vector=chunk['embedding'],
                payload={
                    "document_id": document_id,
                    "chunk_index": idx,
                    "content": chunk['content'],
                    "content_preview": chunk['content'][:100],
                    "page_number": chunk.get('page_number', 1)
                }
            )
            points.append(point)

        # Upload points to Qdrant
        self.qdrant_client.upsert(
            collection_name=self.collection_name,
            points=points
        )

        return [point.id for point in points]


async def main():
    if len(sys.argv) != 2:
        print("Usage: python standalone_ingest.py <path_to_document_or_directory>")
        sys.exit(1)

    input_path = sys.argv[1]
    if not os.path.exists(input_path):
        print(f"Error: Path {input_path} does not exist")
        sys.exit(1)

    # Get API keys from environment or config
    cohere_api_key = os.getenv("COHERE_API_KEY", "")
    qdrant_host = os.getenv("QDRANT_HOST", "")
    qdrant_api_key = os.getenv("QDRANT_API_KEY", "")
    database_url = os.getenv("DATABASE_URL", "")

    if not all([cohere_api_key, qdrant_host, qdrant_api_key, database_url]):
        print("Error: Please set the required environment variables: COHERE_API_KEY, QDRANT_HOST, QDRANT_API_KEY, DATABASE_URL")
        sys.exit(1)

    # Initialize the document processor
    processor = StandaloneDocumentProcessor(cohere_api_key, qdrant_host, qdrant_api_key)

    try:
        # Connect to database
        conn = await asyncpg.connect(dsn=database_url)
        print("Connected to database")

        # Check if it's a file or directory
        if os.path.isfile(input_path):
            # Process single file
            await process_single_file(conn, processor, input_path)
        elif os.path.isdir(input_path):
            # Process directory
            await process_directory(conn, processor, input_path)
        else:
            print(f"Error: {input_path} is neither a file nor a directory")
            sys.exit(1)

        await conn.close()
        print("Database connection closed")

    except Exception as e:
        print(f"Error during ingestion: {str(e)}")
        sys.exit(1)


async def process_single_file(conn, processor, file_path):
    """Process a single file"""
    print(f"Starting ingestion of document: {file_path}")

    # Create document record in database
    file_size = os.path.getsize(file_path)
    file_ext = Path(file_path).suffix.lower()

    document_data = {
        'filename': os.path.basename(file_path),
        'original_name': os.path.basename(file_path),
        'content_type': f'application/{file_ext[1:]}' if file_ext else 'application/octet-stream',
        'size_bytes': file_size,
        'checksum': hashlib.md5(open(file_path, 'rb').read()).hexdigest(),
        'status': 'processing',
        'metadata': {}
    }

    # Insert document record
    result = await conn.fetchrow("""
        INSERT INTO documents (
            filename, original_name, content_type, size_bytes,
            pages_count, checksum, metadata, status
        ) VALUES ($1, $2, $3, $4, $5, $6, $7, $8)
        RETURNING id
    """,
    document_data['filename'],
    document_data['original_name'],
    document_data['content_type'],
    document_data['size_bytes'],
    document_data.get('pages_count'),
    document_data['checksum'],
    document_data['metadata'],
    document_data['status']
    )

    document_id = str(result['id'])
    print(f"Created document record with ID: {document_id}")

    try:
        # Process the document and generate chunks with embeddings using the file path
        processed_chunks = await processor.process_file(file_path)
        print(f"Processed document into {len(processed_chunks)} chunks")

        # Store embeddings in Qdrant
        point_ids = processor.store_embeddings(document_id, processed_chunks)
        print(f"Stored embeddings in Qdrant")

        # Update document status to completed
        await conn.execute("""
            UPDATE documents SET status = $1, processed_at = NOW() WHERE id = $2
        """, 'completed', document_id)
        print(f"Document {document_id} ingestion completed successfully")

    except Exception as e:
        print(f"Error processing document: {e}")
        await conn.execute("""
            UPDATE documents SET status = $1, processed_at = NOW() WHERE id = $2
        """, 'failed', document_id)
        raise


async def process_directory(conn, processor, directory_path):
    """Process all supported files in a directory"""
    print(f"Starting ingestion of directory: {directory_path}")

    # Supported file extensions
    supported_extensions = {
        '.pdf', '.docx', '.txt', '.md', '.html', '.htm',
        '.py', '.js', '.ts', '.java', '.cpp', '.c',
        '.json', '.xml', '.csv', '.jpg', '.jpeg', '.png', '.bmp'
    }

    # Find all supported files in the directory and subdirectories
    files_to_process = []
    for root, dirs, files in os.walk(directory_path):
        for file in files:
            file_path = os.path.join(root, file)
            if Path(file_path).suffix.lower() in supported_extensions:
                files_to_process.append(file_path)

    print(f"Found {len(files_to_process)} files to process")

    success_count = 0
    failure_count = 0

    for file_path in files_to_process:
        try:
            print(f"Processing file: {file_path}")
            await process_single_file(conn, processor, file_path)
            success_count += 1
        except Exception as e:
            print(f"Error processing {file_path}: {str(e)}")
            failure_count += 1

    print(f"\nIngestion completed. Success: {success_count}, Failed: {failure_count}")


if __name__ == "__main__":
    asyncio.run(main())