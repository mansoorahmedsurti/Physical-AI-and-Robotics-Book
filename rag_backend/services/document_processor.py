import asyncio
import os
import tempfile
from typing import List, Dict, Any, Tuple
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
from io import BytesIO
import json
from ..utils import logger, InvalidDocumentException
from .qdrant_service import qdrant_service
from .embedding_service import EmbeddingService
from ..config import settings

class DocumentProcessor:
    def __init__(self):
        self.embedding_service = EmbeddingService()

    async def process_document(self, filename: str, content: str) -> List[Dict[str, Any]]:
        """
        Process document content and return chunks with embeddings
        Enforce cloud-based Cohere embeddings only - no local fallback allowed
        """
        try:
            # Check content length to prevent memory issues with very large documents
            max_content_length = 500000  # 500k characters
            if len(content) > max_content_length:
                logger.warning(f"Truncating content for {filename} from {len(content)} to {max_content_length} characters")
                content = content[:max_content_length]

            # For now, we'll split the content into chunks
            # In a real implementation, we'd use proper document parsing libraries
            chunks = self._chunk_text(content)

            # Generate embeddings for each chunk using Cohere cloud service
            chunk_contents = [chunk['content'] for chunk in chunks]
            chunk_embeddings = await self.embedding_service.generate_embeddings(chunk_contents)

            # Combine chunks with their embeddings
            processed_chunks = []
            for i, chunk in enumerate(chunks):
                processed_chunk = {
                    'content': chunk['content'],
                    'embedding': chunk_embeddings[i],
                    'page_number': chunk.get('page_number', 1)
                }
                processed_chunks.append(processed_chunk)

            return processed_chunks
        except Exception as e:
            logger.error(f"Critical error processing document {filename}: {str(e)}")
            # According to requirements, no local model fallback is allowed due to 8GB RAM constraint
            # So we must fail completely if embeddings cannot be generated via Cohere
            raise Exception(f"Document processing failed and local fallback is prohibited by requirements: {str(e)}")

    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks
        """
        # Validate parameters to prevent issues
        if chunk_size <= 0:
            chunk_size = 1000
        if overlap >= chunk_size:
            overlap = max(0, chunk_size - 1)  # Ensure overlap is less than chunk_size

        chunks = []

        # If text is empty, return empty list
        if not text:
            return chunks

        start = 0

        while start < len(text):
            # Calculate end position
            end = min(start + chunk_size, len(text))

            # Extract the chunk
            chunk_text = text[start:end]

            chunks.append({
                'content': chunk_text,
                'start_pos': start,
                'end_pos': end
            })

            # Move start position forward, accounting for overlap
            # Ensure we always advance by at least 1 character to prevent infinite loops
            next_start = end - overlap

            # If next_start is not advancing, force advancement
            if next_start <= start:
                next_start = start + 1

            # If we've reached or passed the end of text, break
            if next_start >= len(text):
                break

            start = next_start

            # Additional safety check to prevent creating too many chunks
            if len(chunks) > 10000:  # Reasonable upper limit
                logger.warning(f"Too many chunks created ({len(chunks)}), stopping to prevent memory exhaustion")
                break

        return chunks

    async def process_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process a file from disk and return chunks with embeddings
        Enforce cloud-based Cohere embeddings only - no local fallback allowed
        """
        try:
            # Check file size to prevent memory issues with very large files
            file_size = os.path.getsize(file_path)
            max_size = 10 * 1024 * 1024  # 10 MB limit
            if file_size > max_size:
                raise InvalidDocumentException(f"File too large ({file_size} bytes), maximum allowed size is {max_size} bytes")

            file_extension = Path(file_path).suffix.lower()

            if file_extension == '.pdf':
                content = self._extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                content = self._extract_text_from_docx(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp']:
                content = self._extract_text_from_image(file_path)
            elif file_extension in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv']:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
            else:
                # For other files, try to read as text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()

            # Limit content size to prevent memory issues
            max_content_length = 500000  # 500k characters
            if len(content) > max_content_length:
                logger.warning(f"Truncating content for {file_path} from {len(content)} to {max_content_length} characters")
                content = content[:max_content_length]

            return await self.process_document(Path(file_path).name, content)
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            # Propagate the error upward - no local fallback allowed due to 8GB RAM constraint
            raise InvalidDocumentException(f"Could not process file {file_path}: {str(e)}")

    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from a PDF file
        """
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text += page.extract_text() + "\n"
                    except Exception as e:
                        logger.warning(f"Could not extract text from page {page_num} of {pdf_path}: {str(e)}")
                        # Try alternative method for scanned PDFs using OCR
                        text += self._ocr_pdf_page(pdf_path, page_num) + "\n"
        except Exception as e:
            logger.error(f"Error extracting text from PDF {pdf_path}: {str(e)}")
            raise
        return text

    def _ocr_pdf_page(self, pdf_path: str, page_num: int) -> str:
        """
        Extract text from a PDF page using OCR (for scanned PDFs)
        """
        try:
            # Import pdf2image if available
            from pdf2image import convert_from_path

            # Convert PDF page to image
            pages = convert_from_path(pdf_path, first_page=page_num+1, last_page=page_num+1)

            if pages:
                # Perform OCR on the image
                text = pytesseract.image_to_string(pages[0])
                return text
            else:
                return ""
        except ImportError:
            logger.warning("pdf2image not installed. Install with 'pip install pdf2image' for OCR support.")
            return ""
        except Exception as e:
            logger.warning(f"OCR failed for page {page_num} of {pdf_path}: {str(e)}")
            return ""

    def _extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text from a DOCX file
        """
        try:
            doc = DocxDocument(docx_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)

            # Also extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)

            return '\n'.join(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path}: {str(e)}")
            raise

    def _extract_text_from_image(self, image_path: str) -> str:
        """
        Extract text from an image file using OCR
        """
        try:
            image = Image.open(image_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            logger.error(f"Error extracting text from image {image_path}: {str(e)}")
            raise

# Global document processor instance
document_processor = DocumentProcessor()